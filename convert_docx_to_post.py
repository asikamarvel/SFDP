"""
SFDP Blog Post Converter
========================
Converts .docx files to markdown blog posts for the SFDP website.

Usage:
    python convert_docx_to_post.py "path/to/your/document.docx"
    python convert_docx_to_post.py "path/to/your/document.docx" --title "Override Title" --category "Health"

Requirements:
    pip install python-docx Pillow

Features:
    - Auto-extracts title from first heading in document
    - Auto-detects category based on article keywords
    - Extracts text with formatting (bold, italic, headings)
    - Extracts and saves images
    - Preserves hyperlinks
    - Generates proper folder structure
    - Creates frontmatter automatically
"""

import os
import sys
import re
import argparse
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Please install it by running: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    import io
except ImportError:
    print("WARNING: Pillow is not installed. Image optimization will be skipped.")
    print("Install it by running: pip install Pillow")
    Image = None


def slugify(text):
    """Convert text to URL-friendly slug."""
    # Convert to lowercase
    slug = text.lower()
    # Remove special characters
    slug = re.sub(r'[^\w\s-]', '', slug)
    # Replace spaces with hyphens
    slug = re.sub(r'[\s_]+', '-', slug)
    # Remove multiple hyphens
    slug = re.sub(r'-+', '-', slug)
    # Remove leading/trailing hyphens
    slug = slug.strip('-')
    return slug


# Category keywords for auto-detection
CATEGORY_KEYWORDS = {
    "Disease": [
        "disease", "illness", "disorder", "syndrome", "infection", "fever",
        "malaria", "cholera", "typhoid", "tuberculosis", "diabetes", "cancer",
        "measles", "meningitis", "hepatitis", "pneumonia", "ebola", "lassa",
        "dengue", "yellow fever", "hiv", "aids", "gonorrhea", "gonorrhoea",
        "herpes", "syphilis", "chlamydia", "monkeypox", "smallpox"
    ],
    "Vaccine": [
        "vaccine", "vaccination", "immunization", "immunisation", "inoculation",
        "booster", "dose", "pfizer", "moderna", "astrazeneca", "johnson",
        "mrna", "antibody", "antibodies", "immunize", "immunise"
    ],
    "Mental Health": [
        "mental health", "depression", "anxiety", "stress", "ptsd",
        "psychological", "psychiatry", "psychiatric", "therapy", "counseling",
        "counselling", "bipolar", "schizophrenia", "suicide", "trauma",
        "emotional", "wellness", "mindfulness", "alzheimer"
    ],
    "Virus": [
        "virus", "viral", "covid", "coronavirus", "corona", "sars",
        "influenza", "flu", "pandemic", "epidemic", "outbreak"
    ],
    "Prevention": [
        "prevention", "prevent", "preventive", "protective", "hygiene",
        "sanitation", "wash", "clean", "safe", "safety", "awareness",
        "screening", "checkup", "check-up", "exercise", "diet", "nutrition"
    ],
    "Innovation": [
        "innovation", "technology", "research", "study", "discovery",
        "breakthrough", "treatment", "therapy", "drug", "medicine",
        "clinical trial", "experiment", "science", "scientific"
    ],
    "Charity": [
        "charity", "donation", "donate", "fundraising", "nonprofit",
        "non-profit", "volunteer", "aid", "humanitarian", "relief",
        "foundation", "organization", "organisation", "grant"
    ],
    "News": [
        "news", "update", "announcement", "report", "latest",
        "breaking", "today", "recent", "current"
    ]
}


def detect_category(text):
    """Auto-detect category based on keywords in the text."""
    text_lower = text.lower()
    scores = {}
    
    for category, keywords in CATEGORY_KEYWORDS.items():
        score = 0
        for keyword in keywords:
            # Count occurrences of each keyword
            count = text_lower.count(keyword)
            score += count
        scores[category] = score
    
    # Get the category with highest score
    if scores:
        best_category = max(scores, key=scores.get)
        if scores[best_category] > 0:
            return best_category
    
    # Default to Health if no keywords matched
    return "Health"


def extract_title_from_doc(doc):
    """Extract the title from the first heading or paragraph of the document."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a heading style
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "Title" in style_name:
                return text
            # Otherwise use the first non-empty paragraph as title
            return text
    return None


def get_full_document_text(doc):
    """Get all text from document for category detection."""
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return " ".join(full_text)


def get_hyperlink_url(paragraph):
    """Extract hyperlinks from a paragraph."""
    hyperlinks = {}
    for rel in paragraph.part.rels.values():
        if "hyperlink" in rel.reltype:
            hyperlinks[rel.rId] = rel.target_ref
    return hyperlinks


def extract_paragraph_text_with_formatting(paragraph):
    """Extract text from paragraph with markdown formatting."""
    markdown_text = ""
    hyperlinks = {}
    
    # Get hyperlinks from relationships
    for rel in paragraph.part.rels.values():
        if "hyperlink" in rel.reltype:
            hyperlinks[rel.rId] = rel.target_ref
    
    for child in paragraph._element:
        # Check for hyperlinks
        if child.tag == qn('w:hyperlink'):
            rId = child.get(qn('r:id'))
            link_text = ""
            for run_elem in child.findall(qn('w:r')):
                for text_elem in run_elem.findall(qn('w:t')):
                    if text_elem.text:
                        link_text += text_elem.text
            
            if rId and rId in hyperlinks:
                markdown_text += f"[{link_text}]({hyperlinks[rId]})"
            else:
                markdown_text += link_text
        
        # Check for regular runs
        elif child.tag == qn('w:r'):
            run_text = ""
            is_bold = False
            is_italic = False
            footnote_ref = None
            
            # Check formatting
            rPr = child.find(qn('w:rPr'))
            if rPr is not None:
                if rPr.find(qn('w:b')) is not None:
                    is_bold = True
                if rPr.find(qn('w:i')) is not None:
                    is_italic = True
            
            # Check for footnote/endnote reference
            footnoteRef = child.find(qn('w:footnoteReference'))
            if footnoteRef is not None:
                footnote_ref = footnoteRef.get(qn('w:id'))
            
            endnoteRef = child.find(qn('w:endnoteReference'))
            if endnoteRef is not None:
                footnote_ref = endnoteRef.get(qn('w:id'))
            
            # Get text
            for text_elem in child.findall(qn('w:t')):
                if text_elem.text:
                    run_text += text_elem.text
            
            # Apply markdown formatting
            if run_text:
                if is_bold and is_italic:
                    run_text = f"***{run_text}***"
                elif is_bold:
                    run_text = f"**{run_text}**"
                elif is_italic:
                    run_text = f"*{run_text}*"
                markdown_text += run_text
            
            # Add footnote reference marker
            if footnote_ref:
                markdown_text += f"[^{footnote_ref}]"
    
    return markdown_text


def get_heading_level(paragraph):
    """Determine the heading level of a paragraph."""
    style_name = paragraph.style.name if paragraph.style else ""
    
    if "Heading 1" in style_name or style_name == "Title":
        return 1
    elif "Heading 2" in style_name:
        return 2
    elif "Heading 3" in style_name:
        return 3
    elif "Heading 4" in style_name:
        return 4
    elif "Heading 5" in style_name:
        return 5
    elif "Heading 6" in style_name:
        return 6
    
    return 0


def is_list_paragraph(paragraph):
    """Check if paragraph is a list item (bullet or numbered)."""
    # Check the paragraph's XML for numbering
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            return True
    
    # Also check style name for list styles
    style_name = paragraph.style.name if paragraph.style else ""
    list_styles = ['List', 'Bullet', 'Number', 'list', 'bullet', 'number']
    for ls in list_styles:
        if ls in style_name:
            return True
    
    return False


def get_list_type(paragraph):
    """Determine if list is bulleted or numbered. Returns 'bullet', 'number', or None."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            # Check numId to determine type (usually odd = bullet, even = numbered, but this varies)
            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))
            if numId is not None:
                # Default to bullet since we can't easily determine
                return 'bullet'
    
    style_name = paragraph.style.name if paragraph.style else ""
    if 'Number' in style_name or 'number' in style_name:
        return 'number'
    
    return 'bullet'


def is_reference_section(text):
    """Check if this paragraph starts a references section."""
    text_lower = text.lower().strip()
    reference_headers = [
        'references', 'reference', 'bibliography', 'sources', 'citations',
        'works cited', 'further reading', 'resources', 'links'
    ]
    for header in reference_headers:
        if text_lower == header or text_lower.startswith(header + ':'):
            return True
    return False


def format_callout(textbox_content):
    """Format a textbox as a styled callout box."""
    # Clean up excessive spaces
    textbox_content = re.sub(r'\s+', ' ', textbox_content)
    textbox_content = textbox_content.replace(' \n', '\n')
    
    lines = textbox_content.split('\n')
    
    # Check if it's a quote (starts with quote mark)
    if textbox_content.strip().startswith('"') or textbox_content.strip().startswith('"'):
        # Format as quote callout
        quote_text = textbox_content.strip().strip('"').strip('"').strip()
        return f'\n<div class="quote-callout">\n{quote_text}\n</div>\n'
    
    # Otherwise, format as regular callout box
    if lines:
        title = lines[0].strip()
        content_lines = lines[1:] if len(lines) > 1 else []
        content = '\n'.join(content_lines).strip()
        
        # Skip video placeholders
        if '[Featured' in title or 'video' in title.lower():
            return ''
        
        html = '\n<div class="callout-box">\n'
        html += f'<h3>{title}</h3>\n'
        if content:
            # Split into paragraphs
            paragraphs = content.split('\n')
            for p in paragraphs:
                if p.strip():
                    html += f'<p>{p.strip()}</p>\n'
        html += '</div>\n'
        return html
    
    return ''


def extract_textboxes_with_positions(doc):
    """Extract text from text boxes/shapes with their positions in the document."""
    textboxes = []
    seen_content = set()  # Track seen content to avoid duplicates
    
    try:
        body = doc._element.body
        para_count = 0
        
        for elem in body:
            if elem.tag.endswith('}p'):  # Paragraph
                para_count += 1
            
            # Look for textboxes in this element
            for txbx in elem.findall('.//w:txbxContent', 
                namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                text_parts = []
                
                # Extract all paragraphs in the textbox
                for para in txbx.findall('.//w:p',
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    para_text = []
                    for t_elem in para.findall('.//w:t',
                        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if t_elem.text:
                            para_text.append(t_elem.text)
                    if para_text:
                        text_parts.append(' '.join(para_text))
                
                if text_parts:
                    textbox_content = '\n'.join(text_parts).strip()
                    # Check if we've seen this content before (avoid duplicates)
                    if textbox_content and textbox_content not in seen_content:
                        seen_content.add(textbox_content)
                        textboxes.append({
                            'position': para_count,
                            'content': textbox_content
                        })
    
    except Exception as e:
        pass  # No textboxes or error accessing them
    
    return textboxes


def extract_footnotes_from_docx(doc):
    """Extract footnotes and endnotes from the document."""
    footnotes = {}
    
    try:
        # Try to access footnotes
        if hasattr(doc, '_part') and hasattr(doc._part, 'footnotes_part'):
            footnotes_part = doc._part.footnotes_part
            if footnotes_part:
                for footnote in footnotes_part.element.findall('.//w:footnote', 
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    footnote_id = footnote.get(qn('w:id'))
                    if footnote_id and footnote_id not in ['0', '-1']:  # Skip separator and continuation footnotes
                        # Extract text from footnote
                        text_parts = []
                        for t_elem in footnote.findall('.//w:t', 
                            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            if t_elem.text:
                                text_parts.append(t_elem.text)
                        if text_parts:
                            footnotes[footnote_id] = ' '.join(text_parts).strip()
    except Exception as e:
        pass  # No footnotes or error accessing them
    
    # Try to access endnotes
    try:
        if hasattr(doc, '_part') and hasattr(doc._part, 'endnotes_part'):
            endnotes_part = doc._part.endnotes_part
            if endnotes_part:
                for endnote in endnotes_part.element.findall('.//w:endnote',
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    endnote_id = endnote.get(qn('w:id'))
                    if endnote_id and endnote_id not in ['0', '-1']:
                        text_parts = []
                        for t_elem in endnote.findall('.//w:t',
                            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            if t_elem.text:
                                text_parts.append(t_elem.text)
                        if text_parts:
                            footnotes[endnote_id] = ' '.join(text_parts).strip()
    except Exception as e:
        pass  # No endnotes or error accessing them
    
    return footnotes


def extract_images_from_docx(doc, output_folder):
    """Extract all images from the document."""
    images = []
    image_counter = 1
    
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            
            # Determine image extension
            content_type = rel.target_part.content_type
            if "png" in content_type:
                ext = "png"
            elif "jpeg" in content_type or "jpg" in content_type:
                ext = "jpg"
            elif "gif" in content_type:
                ext = "gif"
            elif "webp" in content_type:
                ext = "webp"
            else:
                ext = "png"  # Default
            
            # Generate filename
            filename = f"image_{image_counter:03d}.{ext}"
            filepath = os.path.join(output_folder, filename)
            
            # Save image
            with open(filepath, 'wb') as f:
                f.write(image_data)
            
            # Optimize image if Pillow is available
            if Image:
                try:
                    img = Image.open(filepath)
                    # Convert to RGB if necessary
                    if img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')
                        filepath = filepath.rsplit('.', 1)[0] + '.jpg'
                        filename = filename.rsplit('.', 1)[0] + '.jpg'
                    # Resize if too large
                    max_width = 1200
                    if img.width > max_width:
                        ratio = max_width / img.width
                        new_height = int(img.height * ratio)
                        img = img.resize((max_width, new_height), Image.LANCZOS)
                    img.save(filepath, quality=85, optimize=True)
                except Exception as e:
                    print(f"  Warning: Could not optimize {filename}: {e}")
            
            images.append({
                'rId': rel.rId,
                'filename': filename,
                'filepath': filepath
            })
            image_counter += 1
    
    return images


def find_image_in_paragraph(paragraph, images):
    """Check if paragraph contains an image and return its filename."""
    for run in paragraph.runs:
        for elem in run._element:
            # Look for drawings (images)
            drawings = elem.findall('.//' + qn('a:blip'), 
                                   namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            for drawing in drawings:
                embed = drawing.get(qn('r:embed'))
                if embed:
                    for img in images:
                        if img['rId'] == embed:
                            return img['filename']
    
    # Alternative method - check for inline shapes
    for child in paragraph._element.iter():
        if 'blip' in child.tag:
            embed = child.get(qn('r:embed'))
            if embed:
                for img in images:
                    if img['rId'] == embed:
                        return img['filename']
    
    return None


def convert_corrupted_docx(docx_path, title=None, category=None, date=None):
    """Convert a corrupted .docx file by extracting text directly from XML."""
    import zipfile
    
    print("Attempting to recover text from corrupted document...")
    
    with zipfile.ZipFile(docx_path, 'r') as zf:
        try:
            xml_content = zf.read('word/document.xml').decode('utf-8')
        except:
            # Try reading with CRC check disabled
            for info in zf.infolist():
                if info.filename == 'word/document.xml':
                    with zf.open(info) as f:
                        xml_content = f.read().decode('utf-8')
                    break
    
    # Extract text from XML
    texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
    
    if not texts:
        print("ERROR: Could not extract any text from document")
        return False, None, None
    
    # First text is likely the title
    if not title:
        title = texts[0].strip()
        print(f"📌 Auto-extracted title: {title}")
    
    # Detect category from all text
    full_text = ' '.join(texts)
    if not category:
        category = detect_category(full_text)
        print(f"📂 Auto-detected category: {category}")
    
    # Generate slug and paths
    slug = slugify(title)
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")
    
    # Get the posts directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    posts_dir = os.path.join(script_dir, "posts")
    post_dir = os.path.join(posts_dir, slug)
    images_dir = os.path.join(post_dir, "images")
    
    # Create directories
    os.makedirs(images_dir, exist_ok=True)
    print(f"Created post directory: {post_dir}")
    
    # Build markdown content (simple paragraphs since we lost structure)
    markdown_content = []
    current_para = []
    
    for text in texts[1:]:  # Skip title
        text = text.strip()
        if text:
            current_para.append(text)
        else:
            if current_para:
                markdown_content.append(' '.join(current_para))
                current_para = []
    
    if current_para:
        markdown_content.append(' '.join(current_para))
    
    # Create frontmatter
    frontmatter = f'''---
title: "{title}"
date: {date}
category: "{category}"
---
'''
    
    # Combine content
    full_content = frontmatter + "\n\n".join(markdown_content)
    
    # Clean up
    full_content = re.sub(r'\n{3,}', '\n\n', full_content)
    
    # Write markdown file
    md_path = os.path.join(post_dir, "index.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(full_content)
    
    print(f"\n⚠️  Document was corrupted - images could not be extracted")
    print(f"✅ Successfully created blog post (text only)!")
    print(f"   Location: {post_dir}")
    print(f"   Markdown: {md_path}")
    print(f"\n📝 Post URL will be: blog-post.html?post={slug}")
    
    # Update blog index
    update_blog_index(slug, title, date, category)
    
    return True, title, category


def convert_docx_to_markdown(docx_path, title=None, category=None, date=None):
    """Convert a .docx file to markdown and create the blog post structure."""
    import zipfile
    
    # Validate input file
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}")
        return False, None, None
    
    # Load document
    print(f"Loading document: {docx_path}")
    try:
        doc = Document(docx_path)
    except zipfile.BadZipFile as e:
        print(f"WARNING: Document has corrupted images, attempting text-only recovery...")
        # Fall back to direct XML extraction for corrupted files
        try:
            return convert_corrupted_docx(docx_path, title, category, date)
        except Exception as e2:
            print(f"ERROR: Could not recover document: {e2}")
            return False, None, None
    except Exception as e:
        print(f"ERROR: Could not read document: {e}")
        return False, None, None
    
    # Auto-extract title if not provided
    if not title:
        title = extract_title_from_doc(doc)
        if not title:
            print("ERROR: Could not extract title from document.")
            return False, None, None
        print(f"📌 Auto-extracted title: {title}")
    
    # Auto-detect category if not provided
    if not category:
        full_text = get_full_document_text(doc)
        category = detect_category(full_text)
        print(f"📂 Auto-detected category: {category}")
    
    # Generate slug and paths
    slug = slugify(title)
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")
    
    # Get the posts directory (same folder as this script)
    script_dir = os.path.dirname(os.path.abspath(__file__))
    posts_dir = os.path.join(script_dir, "posts")
    post_dir = os.path.join(posts_dir, slug)
    images_dir = os.path.join(post_dir, "images")
    
    # Create directories
    os.makedirs(images_dir, exist_ok=True)
    print(f"Created post directory: {post_dir}")
    
    # Extract images
    print("Extracting images...")
    images = extract_images_from_docx(doc, images_dir)
    print(f"  Found {len(images)} images")
    
    # Extract footnotes
    print("Extracting footnotes...")
    footnotes = extract_footnotes_from_docx(doc)
    print(f"  Found {len(footnotes)} footnotes/endnotes")
    
    # Extract textboxes/sidebars with positions
    print("Extracting textboxes/sidebars...")
    textboxes_with_pos = extract_textboxes_with_positions(doc)
    print(f"  Found {len(textboxes_with_pos)} textboxes/sidebars")
    
    # Create a map of textboxes by position
    textbox_map = {tb['position']: tb['content'] for tb in textboxes_with_pos}
    
    # Convert document to markdown
    print("Converting content to markdown...")
    markdown_content = []
    cover_image = None
    image_index = 0
    first_text_skipped = False  # Flag to skip the title (first heading/paragraph)
    in_list = False  # Track if we're in a list
    list_counter = 0  # For numbered lists
    in_references = False  # Track if we're in references section
    para_index = 0  # Track paragraph position for textbox insertion
    
    for para in doc.paragraphs:
        para_index += 1
        
        # Skip empty paragraphs
        if not para.text.strip() and not find_image_in_paragraph(para, images):
            # End list if we hit an empty paragraph
            if in_list:
                in_list = False
                list_counter = 0
            # Check if there's a textbox after this position
            if para_index in textbox_map:
                textbox_content = textbox_map[para_index]
                formatted_callout = format_callout(textbox_content)
                markdown_content.append(formatted_callout)
            continue
        
        # Check for images in paragraph
        img_filename = find_image_in_paragraph(para, images)
        if img_filename:
            # First image becomes cover image
            if cover_image is None:
                cover_image = img_filename
            markdown_content.append(f"\n![](images/{img_filename})\n")
            in_list = False
            list_counter = 0
            continue
        
        # Get heading level
        heading_level = get_heading_level(para)
        
        # Extract text with formatting
        text = extract_paragraph_text_with_formatting(para)
        
        if not text.strip():
            continue
        
        # Skip the first text element (title) since it's already in frontmatter
        if not first_text_skipped:
            first_text_skipped = True
            # Only skip if it looks like the title (matches or is a heading)
            if heading_level > 0 or text.strip().lower() == title.lower():
                continue
        
        # Check if this is a references section header
        if is_reference_section(text) or (heading_level > 0 and is_reference_section(para.text)):
            in_references = True
            # Add extra spacing before references
            markdown_content.append("\n\n---\n")
            if heading_level > 0:
                prefix = "#" * heading_level
                markdown_content.append(f"\n{prefix} {text.strip()}\n\n")
            else:
                markdown_content.append(f"\n## References\n\n")
            continue
        
        # Check if this is a list item
        is_list = is_list_paragraph(para)
        
        # Format as heading, list item, or paragraph
        if heading_level > 0:
            in_list = False
            list_counter = 0
            prefix = "#" * heading_level
            markdown_content.append(f"\n{prefix} {text.strip()}\n")
        elif is_list:
            list_type = get_list_type(para)
            if not in_list:
                # Start of a new list - add a blank line before
                markdown_content.append("")
            in_list = True
            if list_type == 'number':
                list_counter += 1
                markdown_content.append(f"{list_counter}. {text.strip()}")
            else:
                markdown_content.append(f"- {text.strip()}")
        elif in_references:
            # In references section, treat each paragraph as a bullet point
            markdown_content.append(f"- {text.strip()}")
        else:
            # Regular paragraph
            if in_list:
                # End of list, add spacing
                in_list = False
                list_counter = 0
                markdown_content.append("")
            markdown_content.append(f"\n{text.strip()}\n")
        
        # Check if there's a textbox after this paragraph
        if para_index in textbox_map:
            textbox_content = textbox_map[para_index]
            formatted_callout = format_callout(textbox_content)
            if formatted_callout:
                markdown_content.append(formatted_callout)
    
    # If no cover image found, check if there are any images
    if cover_image is None and images:
        cover_image = images[0]['filename']
    
    # Create frontmatter
    frontmatter = f'''---
title: "{title}"
date: {date}
category: "{category}"'''
    
    if cover_image:
        frontmatter += f'\ncoverImage: "{cover_image}"'
    
    frontmatter += "\n---\n"
    
    # Combine frontmatter and content
    full_content = frontmatter + "\n".join(markdown_content)
    
    # Add footnotes at the end if any exist
    if footnotes:
        full_content += "\n\n---\n\n## Footnotes\n\n"
        for footnote_id in sorted(footnotes.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            full_content += f"[^{footnote_id}]: {footnotes[footnote_id]}\n\n"
    
    # Clean up multiple blank lines
    full_content = re.sub(r'\n{3,}', '\n\n', full_content)
    
    # Write markdown file
    md_path = os.path.join(post_dir, "index.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(full_content)
    
    print(f"\n✅ Successfully created blog post!")
    print(f"   Location: {post_dir}")
    print(f"   Markdown: {md_path}")
    print(f"   Images: {len(images)} files in {images_dir}")
    if footnotes:
        print(f"   Footnotes: {len(footnotes)} preserved")
    if textboxes_with_pos:
        print(f"   Sidebars: {len(textboxes_with_pos)} inserted contextually")
    print(f"\n📝 Post URL will be: blog-post.html?post={slug}")
    
    # Update blog index (optional)
    update_blog_index(slug, title, date, category)
    
    return True, title, category


def update_blog_index(slug, title, date, category):
    """Update the blog-index.json file with the new post."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    index_path = os.path.join(script_dir, "blog", "blog-index.json")
    
    try:
        import json
        
        # Load existing index (handle various encodings)
        if os.path.exists(index_path):
            # Try utf-8-sig first (handles BOM), then utf-8
            for encoding in ['utf-8-sig', 'utf-8', 'utf-16']:
                try:
                    with open(index_path, 'r', encoding=encoding) as f:
                        index = json.load(f)
                    break
                except (UnicodeDecodeError, json.JSONDecodeError):
                    continue
            else:
                print(f"   Warning: Could not read blog-index.json, creating new one")
                index = []
        else:
            index = []
        
        # Check if post already exists
        for post in index:
            if post.get('slug') == slug:
                print(f"   Note: Post '{slug}' already exists in blog-index.json")
                return
        
        # Add new post
        new_post = {
            "slug": slug,
            "title": title,
            "date": date,
            "category": category
        }
        
        # Add the new post
        index.append(new_post)
        
        # Sort by date (newest first)
        index.sort(key=lambda x: x.get('date', '0000-00-00'), reverse=True)
        
        # Save updated index (always use utf-8 without BOM)
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(index, f, indent=2, ensure_ascii=False)
        
        print(f"   Updated blog-index.json (sorted by date)")
        
    except Exception as e:
        print(f"   Warning: Could not update blog-index.json: {e}")
    
    # Also update editorials.html
    update_editorials_page(slug, title, date, category)
    
    # Also update homepage with latest 3 posts
    update_homepage(slug, title, date, category)


def update_editorials_page(slug, title, date, category):
    """Update the editorials.html page with the new post."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    editorials_path = os.path.join(script_dir, "editorials.html")
    
    try:
        if not os.path.exists(editorials_path):
            print(f"   Warning: editorials.html not found")
            return
        
        with open(editorials_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create the new post entry
        # Escape single quotes in title for JavaScript
        escaped_title = title.replace("'", "\\'")
        new_entry = f"{{ slug: '{slug}', title: '{escaped_title}', date: '{date}', category: '{category}' }},"
        
        # Find where to insert (after "const allPosts = [")
        marker = "const allPosts = [\n"
        if marker in content:
            insert_pos = content.find(marker) + len(marker)
            # Add proper indentation
            new_line = f"      {new_entry}\n"
            
            # Check if post already exists
            if slug in content:
                print(f"   Note: Post already exists in editorials.html")
                return
            
            # Insert the new post
            content = content[:insert_pos] + new_line + content[insert_pos:]
            
            with open(editorials_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"   Updated editorials.html")
        else:
            print(f"   Warning: Could not find allPosts array in editorials.html")
    
    except Exception as e:
        print(f"   Warning: Could not update editorials.html: {e}")


def update_homepage(slug, title, date, category, excerpt=""):
    """Update the index.html homepage with the latest 3 blog posts."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    index_path = os.path.join(script_dir, "index.html")
    blog_index_path = os.path.join(script_dir, "blog", "blog-index.json")
    
    try:
        if not os.path.exists(index_path):
            print(f"   Warning: index.html not found")
            return
        
        # Load blog index to get latest 3 posts
        posts = []
        if os.path.exists(blog_index_path):
            with open(blog_index_path, 'r', encoding='utf-8') as f:
                posts = json.load(f)
        
        # Sort by date to ensure latest first
        posts.sort(key=lambda x: x.get('date', '0000-00-00'), reverse=True)
        
        # Get top 3 posts
        latest_posts = posts[:3] if len(posts) >= 3 else posts
        
        if not latest_posts:
            print(f"   Warning: No posts found in blog-index.json")
            return
        
        with open(index_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Find the posts-grid section for editorials
        start_marker = '<div class="posts-grid">'
        end_marker = '</div>\n    </div>\n  </section>\n\n  <!-- ==================== CTA Section ==================== -->'
        
        start_pos = content.find(start_marker)
        end_pos = content.find(end_marker)
        
        if start_pos == -1 or end_pos == -1:
            print(f"   Warning: Could not find editorials section in index.html")
            return
        
        # Build the new HTML for the 3 cards
        cards_html = f'      {start_marker}\n'
        
        for i, post in enumerate(latest_posts):
            post_slug = post['slug']
            post_title = post['title']
            post_date = post['date']
            post_category = post['category']
            
            # Format date nicely
            try:
                date_obj = datetime.strptime(post_date, '%Y-%m-%d')
                formatted_date = date_obj.strftime('%B %d, %Y')
            except:
                formatted_date = post_date
            
            # Generate placeholder image based on category
            category_images = {
                'Health': 'https://images.unsplash.com/photo-1559757175-5700dde675bc?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80',
                'Disease': 'https://images.unsplash.com/photo-1584515933487-779824d29309?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80',
                'Vaccine': 'https://images.unsplash.com/photo-1587854692152-cbe660dbde88?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80',
                'Prevention': 'https://images.unsplash.com/photo-1505751172876-fa1923c5c528?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80',
                'Mental Health': 'https://images.unsplash.com/photo-1573496359142-b8d87734a5a2?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80',
                'Nutrition': 'https://images.unsplash.com/photo-1492725764893-90b379c2b6e7?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80',
                'Cancer Awareness': 'https://images.unsplash.com/photo-1579154204601-01588f351e67?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80'
            }
            image_url = category_images.get(post_category, 'https://images.unsplash.com/photo-1576091160399-112ba8d25d1d?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80')
            
            # Create excerpt from title or use provided excerpt
            excerpt_text = post_title[:120] + "..." if len(post_title) > 120 else post_title
            
            delay_class = f' delay-{i}' if i > 0 else ''
            
            cards_html += f'''        <article class="card blog-card scroll-animate{delay_class}">
          <img src="{image_url}" alt="{post_title[:50]}" class="card-img">
          <div class="card-body">
            <span class="blog-category">{post_category}</span>
            <h3 class="card-title">
              <a href="blog-post.html?post={post_slug}">{post_title}</a>
            </h3>
            <p class="card-text">{excerpt_text}</p>
            <div class="card-meta">
              <i class="far fa-calendar-alt"></i> {formatted_date}
            </div>
          </div>
        </article>
        
'''
        
        cards_html += '      </div>'
        
        # Replace the section
        content = content[:start_pos] + cards_html + content[end_pos:]
        
        with open(index_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"   Updated index.html with latest 3 posts")
        
    except Exception as e:
        print(f"   Warning: Could not update index.html: {e}")


def analyze_document(docx_path):
    """Analyze document and return suggested title and category without converting."""
    import zipfile
    
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}")
        return None, None
    
    try:
        doc = Document(docx_path)
    except zipfile.BadZipFile as e:
        # Try to open with strict=False to handle minor corruption
        print(f"WARNING: Document may be slightly corrupted, attempting recovery...", file=sys.stderr)
        try:
            # Read the document text directly from the XML
            import zipfile
            with zipfile.ZipFile(docx_path, 'r') as zf:
                # Set strict to False to ignore CRC errors
                zf._strict = False
                try:
                    xml_content = zf.read('word/document.xml').decode('utf-8')
                    # Extract text between <w:t> tags
                    import re
                    texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                    if texts:
                        title = texts[0].strip() if texts else None
                        full_text = ' '.join(texts)
                        category = detect_category(full_text)
                        return title, category
                except:
                    pass
        except:
            pass
        print(f"ERROR: Could not read document: {e}")
        return None, None
    except Exception as e:
        print(f"ERROR: Could not read document: {e}")
        return None, None
    
    # Extract title
    title = extract_title_from_doc(doc)
    
    # Detect category
    full_text = get_full_document_text(doc)
    category = detect_category(full_text)
    
    return title, category


def main():
    parser = argparse.ArgumentParser(
        description="Convert a .docx file to a markdown blog post for SFDP website.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python convert_docx_to_post.py "article.docx"
  python convert_docx_to_post.py "article.docx" --title "My Health Article" --category "Health"
  python convert_docx_to_post.py "article.docx" --analyze

Categories: Health, Disease, Vaccine, Prevention, Mental Health, Virus, Innovation, News, Charity
        """
    )
    
    parser.add_argument("docx_file", help="Path to the .docx file to convert")
    parser.add_argument("-t", "--title", default=None, 
                       help="Title of the blog post (auto-extracted if not provided)")
    parser.add_argument("-c", "--category", default=None, 
                       help="Category (auto-detected if not provided)")
    parser.add_argument("-d", "--date", default=None,
                       help="Publication date (YYYY-MM-DD, default: today)")
    parser.add_argument("--analyze", action="store_true",
                       help="Only analyze and show suggested title/category without converting")
    
    args = parser.parse_args()
    
    # Analyze-only mode
    if args.analyze:
        title, category = analyze_document(args.docx_file)
        if title:
            print(f"TITLE:{title}")
            print(f"CATEGORY:{category}")
            return 0
        return 1
    
    # Convert the document
    success, final_title, final_category = convert_docx_to_markdown(
        args.docx_file,
        args.title,
        args.category,
        args.date
    )
    
    if success:
        print("\n" + "="*50)
        print("NEXT STEPS:")
        print("="*50)
        print("1. Review the generated markdown file")
        print("2. Check that images are properly placed")
        print("3. Test the post by visiting:")
        print(f"   blog-post.html?post={slugify(final_title)}")
        print("="*50)
    
    return 0 if success else 1


if __name__ == "__main__":
    sys.exit(main())
